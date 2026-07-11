"""
:class:`PyXtabond2Results`: formatted summary display and LaTeX/Word export
of GMM estimation output.
"""

import re

import pandas as pd
import numpy as np
import scipy.stats as stats
from datetime import datetime
from statsmodels.iolib.summary import Summary


_LATEX_SPECIAL_CHARS = {
    '\\': r'\textbackslash{}',
    '&': r'\&',
    '%': r'\%',
    '$': r'\$',
    '#': r'\#',
    '_': r'\_',
    '{': r'\{',
    '}': r'\}',
    '~': r'\textasciitilde{}',
    '^': r'\textasciicircum{}',
}
_LATEX_SPECIAL_RE = re.compile('|'.join(re.escape(c) for c in _LATEX_SPECIAL_CHARS))


def _escape_latex(text) -> str:
    """Escape LaTeX special characters in a free-text label (variable name,
    diagnostic-test name) before inserting it into a raw LaTeX table row.

    Without this, a name containing one of these characters -- most
    commonly an underscore, as in Stata-style ``_cons`` or this package's
    own ``L1_``/``D_``-prefixed variables -- breaks compilation (LaTeX reads
    a bare ``_`` outside math mode as a syntax error). Only applied on the
    LaTeX export path: ``to_word()`` reuses the same variable/test-name
    strings unescaped, which is correct there.
    """
    return _LATEX_SPECIAL_RE.sub(lambda m: _LATEX_SPECIAL_CHARS[m.group()], str(text))


class PyXtabond2Results:
    """
    A class to store, format, and display the results of a GMM estimation.

    This class handles the presentation of coefficients, standard errors,
    and diagnostic tests (Arellano-Bond AR, Sargan/Hansen, Difference-in-Sargan).
    It also provides export functionalities to LaTeX and Word.
    """
    def __init__(self, beta, se, engine, diag, x_names, dep_var, model_name, step_name,
                 gmm_vars=None, iv_vars=None, lag_limits=None, collapse=False, orthogonal=False, id_col=str, time_col=str):
        self.beta = beta
        self.se = se
        self.engine = engine
        self.diag = diag
        self.x_names = x_names
        self.dep_var = dep_var
        self.small = engine.small
        self.model_name = model_name
        self.step_name = step_name

        self.gmm_vars = gmm_vars or []
        self.iv_vars = iv_vars or []
        self.lag_limits = lag_limits or (1, None)
        self.collapse = collapse
        self.orthogonal = orthogonal

        self.id_col = id_col
        self.time_col = time_col

        self.t_stats = self.beta / self.se

        if self.small:
            is_robust_or_2step = (step_name == "two-step") or hasattr(engine, 'V2_robust')
            self.df_resid = (engine.N_clusters - 1) if is_robust_or_2step else (engine.n_obs - engine.k_vars)

            self.p_values = 2.0 * (1.0 - stats.t.cdf(np.abs(self.t_stats), self.df_resid))
            self.ci_crit = stats.t.ppf(0.975, self.df_resid)
        else:
            self.p_values = 2.0 * (1.0 - stats.norm.cdf(np.abs(self.t_stats)))
            self.ci_crit = stats.norm.ppf(0.975)

        self.ci_lower = self.beta - self.ci_crit * self.se
        self.ci_upper = self.beta + self.ci_crit * self.se

    def summary(self):
        """
        Displays the estimation results with clear, aligned typography.
        Replicates the visual structure of Stata's output.
        """
        smry = Summary()

        # --- PREPARE GLOBAL STATISTICS ---
        if self.model_name == "system GMM":
            n_obs = int(np.sum(self.engine.is_level))
        else:
            n_obs = self.engine.n_obs

        n_groups = self.engine.N_groups
        n_instr = self.engine.n_instruments
        wald_stat, wald_df, wald_p, wald_type = self.diag['wald']
        is_robust_or_2step = (self.step_name == "two-step") or hasattr(self.engine, 'V2_robust')

        class MockModel:
            endog_names = self.dep_var
            exog_names = self.x_names

        class MockResults:
            def __init__(self, parent):
                self.model = MockModel()
                self.params = parent.beta.flatten()
                self.bse = parent.se.flatten()
                self.tvalues = parent.t_stats.flatten()
                self.pvalues = parent.p_values.flatten()
                self.ci_l = parent.ci_lower.flatten()
                self.ci_u = parent.ci_upper.flatten()

            def conf_int(self, alpha=0.05):
                return np.column_stack((self.ci_l, self.ci_u))

        mock_res = MockResults(self)

        cov_type = "Standard"
        if hasattr(self.engine, 'V2_robust'):
            cov_type = "Windmeijer" if self.step_name == "two-step" else "Robust"

        gmm_str = ", ".join(self.gmm_vars) if self.gmm_vars else "None"
        iv_str = ", ".join(self.iv_vars) if self.iv_vars else "None"
        lag_max = self.lag_limits[1] if self.lag_limits[1] is not None else 'Max'
        lags_str = f"({self.lag_limits[0]}, {lag_max})"

        now = datetime.now()

        if wald_type == 'F':
            is_const = '_cons' in self.x_names
            if is_robust_or_2step:
                df_resid = self.engine.N_clusters - (1 if is_const else 0)
            else:
                df_resid = self.engine.n_obs - self.engine.k_vars

            stat_name = f'F-statistic ({int(wald_df)}, {int(df_resid)})'
        else:
            stat_name = f'Wald chi2({int(wald_df)})'

        prob_format = f"{wald_p:.4e}" if wald_p < 0.001 else f"{wald_p:.4f}"

        top_left = [
            ('Dep. Variable:', [self.dep_var]),
            ('Model:', [f"GMM ({self.model_name.capitalize()})"]),
            ('Method:', [self.step_name.title()]),
            ('Date:', [now.strftime('%a, %d %b %Y')]),
            ('Time:', [now.strftime('%H:%M:%S')]),
            ('No. Observations:', [str(n_obs)]),
            ('No. Groups:', [str(n_groups)]),
            ('Obs per group:', [f"{n_obs/n_groups:.2f}"]),
            ('No. Instruments:', [str(n_instr)]),
            ('Covariance Type:', [cov_type])
        ]

        top_right = [
            (stat_name + ':', [f"{wald_stat:.2f}"]),
            ('Prob (Stat):', [prob_format]),
            ('Group Variable:', [str(self.id_col)]),
            ('Time variable:', [str(self.time_col)]),
            ('GMM Variables:', [gmm_str]),
            ('IV Variables:', [iv_str]),
            ('Lag Limits:', [lags_str]),
            ('Collapse:', ["Yes" if self.collapse else "No"]),
            ('Transform (FOD):', ["Yes" if self.orthogonal else "No (FD)"]),
            ('Sample Correction:', ["Yes" if self.small else "No"])
        ]

        # 1. Generate ONLY the header first to capture its exact width
        smry.add_table_2cols(mock_res, gleft=top_left, gright=top_right, title="PyXtabond2 Estimation Results")
        header_text = smry.as_text()
        header_w = max(len(line.rstrip()) for line in header_text.split('\n'))

        # 2. Add the coefficients (free to stretch)
        smry.add_table_params(mock_res, xname=self.x_names, yname=self.dep_var, use_t=self.small)
        final_lines = [line.rstrip() for line in smry.as_text().split('\n')]

        # --- PREPARE RAW DIAGNOSTIC DATA ---
        diag_data = [["Diagnostic Test", "Statistic", "P-value"]]
        for lag, (ar_stat, ar_p) in self.diag['ar'].items():
            diag_data.append([f"Arellano-Bond AR({lag})", f"z = {ar_stat:.2f}", f"Pr > z = {ar_p:.3f}"])
        diag_data.append(["Sargan test of overid.", f"chi2({int(self.diag['df'])}) = {self.diag['sargan'][0]:.2f}", f"Pr > chi2 = {self.diag['sargan'][1]:.3f}"])

        if not np.isnan(self.diag['hansen'][0]):
            diag_data.append(["Hansen test of overid.", f"chi2({int(self.diag['df'])}) = {self.diag['hansen'][0]:.2f}", f"Pr > chi2 = {self.diag['hansen'][1]:.3f}"])
        elif hasattr(self.engine, 'hansen_1step_robust'):
            h_stat, h_p = self.engine.hansen_1step_robust
            diag_data.append(["Hansen test of overid.", f"chi2({int(self.diag['df'])}) = {h_stat:.2f}", f"Pr > chi2 = {h_p:.3f}"])

        diff_tests = self.diag.get('diff_sargan', [])
        diff_data = [["Test", "Statistic", "P-value"]]
        if diff_tests:
            for res in diff_tests:
                diff_data.append([f"{res['name']} (Excl. group)", f"chi2({int(res['df_rest'])}) = {res['stat_rest']:.2f}", f"{res['p_rest']:.3f}"])
                diff_data.append(["  Difference", f"chi2({int(res['diff_df'])}) = {res['diff_stat']:.2f}", f"{res['diff_p']:.3f}"])

        # --- ALIGNMENT LOGIC WITH GENEROUS SPACING ---
        def build_matched_table(data_rows, title, target_w):
            c3_max = max(len(str(r[2])) for r in data_rows)
            c2_max = max(len(str(r[1])) for r in data_rows)

            col3_w = max(20, c3_max + 1)
            col2_w = max(20, c2_max + 2)

            col1_w = target_w - col2_w - col3_w

            c1_max = max(len(str(r[0])) for r in data_rows)
            if col1_w < c1_max:
                col1_w = c1_max
                rem = target_w - col1_w
                col3_w = max(10, rem // 2)
                col2_w = rem - col3_w

            lines = []
            lines.append("")
            lines.append(title.center(target_w))
            lines.append('=' * target_w)

            lines.append(f"{str(data_rows[0][0]):<{col1_w}}{str(data_rows[0][1]):>{col2_w}}{str(data_rows[0][2]):>{col3_w}}")
            lines.append('-' * target_w)

            for row in data_rows[1:]:
                lines.append(f"{str(row[0]):<{col1_w}}{str(row[1]):>{col2_w}}{str(row[2]):>{col3_w}}")

            return lines

        # 3. Inject tables
        final_lines.extend(build_matched_table(diag_data, "Diagnostic Tests", header_w))
        final_lines.append('=' * header_w)

        if diff_tests:
            test_type = "Hansen" if hasattr(self.engine, 'V2_robust') or self.step_name == "two-step" else "Sargan"
            final_lines.extend(build_matched_table(diff_data, f"Difference-in-{test_type} Tests", header_w))
            final_lines.append('=' * header_w)

        # Terminal output
        print("\n".join(final_lines))

    # ==========================================
    # --- EXPORT METHODS (LATEX & WORD) ---
    # ==========================================

    def _get_results_df(self):
        """Utility method to structure coefficients in a DataFrame."""
        stat_name = 't' if self.small else 'z'
        df = pd.DataFrame({
            'Coefficient': self.beta.flatten(),
            'Std. Error': self.se.flatten(),
            stat_name: self.t_stats.flatten(),
            f'P>|{stat_name}|': self.p_values.flatten(),
            'CI Lower': self.ci_lower.flatten(),
            'CI Upper': self.ci_upper.flatten()
        }, index=self.x_names)
        return df

    def _get_diagnostics_df(self):
        """Utility method to structure diagnostic tests in a DataFrame."""
        diag_data = [
            [f"Arellano-Bond AR({lag})", f"{ar_stat:.2f}", f"{ar_p:.3f}"]
            for lag, (ar_stat, ar_p) in self.diag['ar'].items()
        ]
        diag_data.append(["Sargan test", f"{self.diag['sargan'][0]:.2f}", f"{self.diag['sargan'][1]:.3f}"])

        if not np.isnan(self.diag['hansen'][0]):
            diag_data.append(["Hansen test", f"{self.diag['hansen'][0]:.2f}", f"{self.diag['hansen'][1]:.3f}"])
        elif hasattr(self.engine, 'hansen_1step_robust'):
            h_stat, h_p = self.engine.hansen_1step_robust
            diag_data.append(["Hansen test", f"{h_stat:.2f}", f"{h_p:.3f}"])

        diff_tests = self.diag.get('diff_sargan', [])
        if diff_tests:
            for res in diff_tests:
                diag_data.append([f"Diff-in-Sargan: {res['name']}", f"{res['diff_stat']:.2f}", f"{res['diff_p']:.3f}"])

        return pd.DataFrame(diag_data, columns=["Test", "Statistic", "P-value"])

    def to_latex(self, filepath=None, full_output=False):
        """
        Generates a LaTeX export of the estimation results.

        Parameters
        ----------
        filepath : str, optional
            If provided, writes the LaTeX output to this file path.
        full_output : bool, optional
            If True, exports the raw console output in a verbatim environment
            (ideal for appendices). If False, generates a clean academic table
            using booktabs style. Default is False.

        Returns
        -------
        str
            The generated LaTeX string.
        """
        import io
        import sys

        if full_output:
            # --- CONSOLE CAPTURE ---
            capture = io.StringIO()
            sys.stdout = capture  # Redirect print statements to our variable
            self.summary()        # Execute summary() silently
            sys.stdout = sys.__stdout__  # Restore normal output
            full_text = capture.getvalue() # Retrieve intercepted text

            latex_str = "% Exact raw output from PyXtabond2 estimation\n"
            latex_str += "\\begin{verbatim}\n"
            latex_str += full_text
            latex_str += "\\end{verbatim}\n"

        else:
            # --- OPTION 2: ACADEMIC TABLE ---
            # _get_results_df/_get_diagnostics_df are shared with to_word(),
            # where variable/test names must stay unescaped -- so escape
            # only these local copies, specific to the LaTeX path.
            res_df = self._get_results_df().rename(index=_escape_latex)
            diag_df = self._get_diagnostics_df()
            diag_df["Test"] = diag_df["Test"].map(_escape_latex)

            # Table header
            latex_str = "\\begin{table}[htbp]\n\\centering\n\\caption{GMM Estimation Results (" + self.model_name + ")}\n"

            # Formatting with booktabs (requires \usepackage{booktabs} in main.tex)
            try:
                # For Pandas >= 1.3.0
                latex_str += res_df.style.format(precision=4).to_latex(hrules=True)
            except AttributeError:
                # Fallback for older Pandas versions
                latex_str += res_df.to_latex(float_format="%.4f", escape=False)

            latex_str += "\n\\vspace{0.5cm}\n"

            # Adding diagnostic tests
            try:
                 latex_str += diag_df.style.hide(axis="index").to_latex(hrules=True)
            except AttributeError:
                 latex_str += diag_df.to_latex(index=False, escape=False)

            latex_str += "\\end{table}"

        # --- COMMON SAVE LOGIC ---
        if filepath:
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(latex_str)
            mode_str = "raw (verbatim)" if full_output else "academic (booktabs)"
            print(f"LaTeX output {mode_str} successfully saved to: {filepath}")

        return latex_str

    def to_word(self, filepath="xtabond2_results.docx", full_output=False):
        """
        Exports the estimation results to a Word document (.docx).

        Requires the 'python-docx' library.

        Parameters
        ----------
        filepath : str, optional
            The file path where the document will be saved. Default is "xtabond2_results.docx".
        full_output : bool, optional
            If True, exports the raw console output in a monospace font.
            If False, generates minimalist academic tables. Default is False.
        """
        import io
        import sys

        try:
            from docx import Document
            from docx.shared import Pt, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            print("Error: The 'python-docx' module is required.")
            return

        doc = Document()

        if full_output:
            doc.add_heading(f"Estimation Results (Raw Output): {self.model_name.capitalize()} GMM", level=0)

            for section in doc.sections:
                section.left_margin = Cm(1.5)
                section.right_margin = Cm(1.5)

            # --- CONSOLE CAPTURE ---
            capture = io.StringIO()
            sys.stdout = capture
            self.summary()
            sys.stdout = sys.__stdout__
            full_text = capture.getvalue()

            p = doc.add_paragraph()
            run = p.add_run(full_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(8)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        else:
            # --- OPTION 2: ACADEMIC TABLES ---
            doc.add_heading(f"Estimation Results: {self.model_name.capitalize()} GMM", level=0)

            # 1. Global metadata
            if self.model_name == "system GMM":
                n_obs = int(np.sum(self.engine.is_level))
            else:
                n_obs = self.engine.n_obs

            doc.add_paragraph(f"Dependent Variable: {self.dep_var}")
            doc.add_paragraph(f"Method: {self.step_name.title()}")
            doc.add_paragraph(f"Observations: {n_obs}")
            doc.add_paragraph(f"Number of groups: {self.engine.N_groups}")
            doc.add_paragraph(f"Number of instruments: {self.engine.n_instruments}")

            # 2. Coefficients table
            doc.add_heading("Coefficients", level=1)
            res_df = self._get_results_df()

            table = doc.add_table(rows=1, cols=len(res_df.columns) + 1)
            table.style = 'Table Grid'

            # Headers
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Variable"
            for i, col_name in enumerate(res_df.columns):
                hdr_cells[i+1].text = col_name

            # Filling variables
            for index, row in res_df.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(index)
                for i, val in enumerate(row):
                    row_cells[i+1].text = f"{val:.4f}" if isinstance(val, (int, float)) else str(val)

            # 3. Diagnostic table
            doc.add_heading("Diagnostic Tests", level=1)
            diag_df = self._get_diagnostics_df()

            table_diag = doc.add_table(rows=1, cols=len(diag_df.columns))
            table_diag.style = 'Table Grid'

            hdr_cells_diag = table_diag.rows[0].cells
            for i, col_name in enumerate(diag_df.columns):
                hdr_cells_diag[i].text = col_name

            for _, row in diag_df.iterrows():
                row_cells = table_diag.add_row().cells
                for i, val in enumerate(row):
                    row_cells[i].text = str(val)

        # --- COMMON SAVE LOGIC ---
        doc.save(filepath)
        mode_str = "raw" if full_output else "academic"
        print(f"Word output ({mode_str}) successfully saved to: {filepath}")
