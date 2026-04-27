import pandas as pd
import numpy as np

class GMMStargazer:
    """
    A utility class to compare multiple PyXtabond2Results models and export 
    the comparative results in a 'stargazer' style (LaTeX and Word).
    
    This class is highly useful for academic research, allowing users to easily 
    generate side-by-side regression tables to demonstrate model robustness 
    across different specifications.

    Parameters
    ----------
    models : list
        A list of PyXtabond2Results instances to be compared.
    model_names : list, optional
        Custom names for the columns representing each model (e.g., ['Baseline', 'Robustness 1']). 
        If not provided, defaults to ['Model 1', 'Model 2', ...].

    Raises
    ------
    ValueError
        If the number of provided `model_names` does not match the number of `models`.

    Examples
    --------
    >>> from pyxtabond2.api import PyXtabond2
    >>> from pyxtabond2.exporter import GMMStargazer
    >>> 
    >>> # Estimate two different models
    >>> model1 = PyXtabond2(df, id_col='id', time_col='year', dep_var='y', 
    ...                     x_vars=['x1'], gmm_vars=['y'], iv_vars=['x2']).fit()
    >>> model2 = PyXtabond2(df, id_col='id', time_col='year', dep_var='y', 
    ...                     x_vars=['x1', 'x3'], gmm_vars=['y'], iv_vars=['x2'], twostep=True).fit()
    >>> 
    >>> # Compare them
    >>> stargazer = GMMStargazer([model1, model2], model_names=["One-Step", "Two-Step Robust"])
    >>> latex_table = stargazer.to_latex("comparison_table.tex")
    """
    def __init__(self, models, model_names=None):
        self.models = models
        self.model_names = model_names or [f"Model {i+1}" for i in range(len(models))]
        
        if len(self.models) != len(self.model_names):
            raise ValueError("The number of model names must match the number of models provided.")
            
        self.all_vars = self._get_unique_vars()
        self.diff_test_keys = self._get_unique_diff_tests()
        self.results_dict = self._extract_data()

    def _get_unique_vars(self):
        """
        Retrieves all unique variables across all models while preserving 
        their order of appearance.
        
        Returns
        -------
        list
            A list of unique variable names.
        """
        seen = set()
        all_vars = []
        for mod in self.models:
            for var in mod.x_names:
                if var not in seen:
                    seen.add(var)
                    all_vars.append(var)
        return all_vars

    def _get_unique_diff_tests(self):
        """
        Dynamically identifies all existing Difference-in-Hansen/Sargan tests 
        across the provided models.
        
        Returns
        -------
        list
            A list of formatted string keys for the difference tests.
        """
        keys = []
        for mod in self.models:
            diff_tests = mod.diag.get('diff_sargan', [])
            for res in diff_tests:
                # Format the row name for the final table
                k = f"Diff p-val: {res['name']}"
                if k not in keys:
                    keys.append(k)
        return keys

    def _get_stars(self, p_val):
        """
        Determines the statistical significance stars based on the p-value.
        
        Parameters
        ----------
        p_val : float
            The p-value to evaluate.

        Returns
        -------
        str
            Asterisks string representing significance levels (* p<0.1, ** p<0.05, *** p<0.01).
        """
        if pd.isna(p_val): return ""
        if p_val < 0.01: return "***"
        if p_val < 0.05: return "**"
        if p_val < 0.10: return "*"
        return ""

    def _extract_data(self):
        """
        Extracts and formats data (coefficients, stats, diagnostics) from all models 
        into a structured dictionary for easy table generation.
        
        Returns
        -------
        dict
            A nested dictionary containing formatted coefficients, statistics, and diagnostic values.
        """
        data = {
            "coeffs": {var: [] for var in self.all_vars},
            "stats": {var: [] for var in self.all_vars}, # Replaces standard errors with t-stats/z-stats
            "diagnostics": {
                "Observations": [],
                "Number of groups": [],
                "Instruments": [],
                "AR(1) p-value": [],
                "AR(2) p-value": [],
                "Sargan p-value": [],
                "Hansen p-value": []
            }
        }

        # Add dynamic keys for Difference tests
        for k in self.diff_test_keys:
            data["diagnostics"][k] = []

        for mod in self.models:
            # 1. Extract coefficients and z-stats/t-stats
            mod_vars = mod.x_names
            for var in self.all_vars:
                if var in mod_vars:
                    idx = mod_vars.index(var)
                    beta = mod.beta.flatten()[idx]
                    # Direct extraction of the z or t statistic
                    stat_val = mod.t_stats.flatten()[idx] 
                    p_val = mod.p_values.flatten()[idx]
                    
                    stars = self._get_stars(p_val)
                    data["coeffs"][var].append(f"{beta:.4f}{stars}")
                    # Enclose the statistic in parentheses
                    data["stats"][var].append(f"({stat_val:.3f})")
                else:
                    data["coeffs"][var].append("")
                    data["stats"][var].append("")

            # 2. Extract standard diagnostics
            n_obs = int(np.sum(mod.engine.is_level)) if mod.model_name == "system GMM" else mod.engine.n_obs
            data["diagnostics"]["Observations"].append(str(n_obs))
            data["diagnostics"]["Number of groups"].append(str(mod.engine.N_groups))
            data["diagnostics"]["Instruments"].append(str(mod.engine.n_instruments))
            
            data["diagnostics"]["AR(1) p-value"].append(f"{mod.diag['ar1'][1]:.3f}" if not pd.isna(mod.diag['ar1'][1]) else "-")
            data["diagnostics"]["AR(2) p-value"].append(f"{mod.diag['ar2'][1]:.3f}" if not pd.isna(mod.diag['ar2'][1]) else "-")
            data["diagnostics"]["Sargan p-value"].append(f"{mod.diag['sargan'][1]:.3f}" if not pd.isna(mod.diag['sargan'][1]) else "-")
            
            hansen_p = "-"
            if not np.isnan(mod.diag['hansen'][1]):
                hansen_p = f"{mod.diag['hansen'][1]:.3f}"
            elif hasattr(mod.engine, 'hansen_1step_robust'):
                hansen_p = f"{mod.engine.hansen_1step_robust[1]:.3f}"
            data["diagnostics"]["Hansen p-value"].append(hansen_p)

            # 3. Dynamic extraction of Difference-in-Hansen/Sargan tests
            # Create a quick dictionary of results specifically for this model
            mod_diff_tests = {f"Diff p-val: {r['name']}": r['diff_p'] for r in mod.diag.get('diff_sargan', [])}
            
            # Populate data based on global keys
            for k in self.diff_test_keys:
                if k in mod_diff_tests:
                    val = mod_diff_tests[k]
                    data["diagnostics"][k].append(f"{val:.3f}" if not pd.isna(val) else "-")
                else:
                    data["diagnostics"][k].append("-")

        return data

    def to_latex(self, filepath=None):
        """
        Generates clean, academic LaTeX code comparable to 'stargazer' output.

        Parameters
        ----------
        filepath : str, optional
            If provided, saves the generated LaTeX code to this file path.

        Returns
        -------
        str
            The generated LaTeX string containing the regression table.
            
        Examples
        --------
        >>> stargazer = GMMStargazer([model1, model2])
        >>> latex_code = stargazer.to_latex()
        >>> print(latex_code)
        """
        n_models = len(self.models)
        
        lines = [
            "\\begin{table}[!htbp]",
            "\\centering",
            "\\caption{GMM Estimation Results}",
            f"\\begin{{tabular}}{{l*{n_models}{{c}}}}",
            "\\toprule",
            " & " + " & ".join([f"\\textbf{{{name}}}" for name in self.model_names]) + " \\\\",
            "\\midrule"
        ]

        # Add coefficients and Stats (z/t)
        for var in self.all_vars:
            coef_line = f"{var} & " + " & ".join(self.results_dict["coeffs"][var]) + " \\\\"
            stat_line = f" & " + " & ".join(self.results_dict["stats"][var]) + " \\\\"
            lines.append(coef_line)
            lines.append(stat_line)
            lines.append("\\vspace{0.1cm} \\\\")

        lines.append("\\midrule")
        
        # Add diagnostics
        for diag_name, values in self.results_dict["diagnostics"].items():
            lines.append(f"{diag_name} & " + " & ".join(values) + " \\\\")
            
            if diag_name == "Instruments" :
                lines.append("\\midrule")  # Add a line after instruments if they are present

        # Update the table footer to specify z-stat/t-stat
        lines.extend([
            "\\bottomrule",
            "\\multicolumn{" + str(n_models + 1) + "}{l}{\\textit{Note:} $^{*} p<0.1$; $^{**} p<0.05$; $^{***} p<0.01$. The $z$ (or $t$) statistics are reported in parentheses.} \\\\",
            "\\end{tabular}",
            "\\end{table}"
        ])

        latex_str = "\n".join(lines)
        
        if filepath:
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(latex_str)
            print(f"LaTeX table saved to: {filepath}")
            
        return latex_str

    def to_word(self, filepath="stargazer_results.docx"):
        """
        Generates a comparative table in a Microsoft Word document.
        
        Requires the 'python-docx' library.

        Parameters
        ----------
        filepath : str, optional
            The destination file path. Default is "stargazer_results.docx".
            
        Examples
        --------
        >>> stargazer = GMMStargazer([model1, model2])
        >>> stargazer.to_word("Chapter_4_Comparisons.docx")
        """
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            print("Error: The 'python-docx' module is required. Install it using 'pip install python-docx'.")
            return

        doc = Document()
        doc.add_heading("GMM Model Comparison", level=1)

        n_models = len(self.models)
        table = doc.add_table(rows=1, cols=n_models + 1)
        table.style = 'Table Grid'

        # Headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Variable"
        for i, name in enumerate(self.model_names):
            hdr_cells[i+1].text = name
            hdr_cells[i+1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Fill variables (Coeffs then Stats)
        for var in self.all_vars:
            row_coef = table.add_row().cells
            row_coef[0].text = var
            for i, val in enumerate(self.results_dict["coeffs"][var]):
                row_coef[i+1].text = val
                row_coef[i+1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            row_stat = table.add_row().cells
            row_stat[0].text = ""
            for i, val in enumerate(self.results_dict["stats"][var]):
                row_stat[i+1].text = val
                row_stat[i+1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.add_row()

        # Fill diagnostics
        for diag_name, values in self.results_dict["diagnostics"].items():
            row_diag = table.add_row().cells
            row_diag[0].text = diag_name
            for i, val in enumerate(values):
                row_diag[i+1].text = val
                row_diag[i+1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add legend
        doc.add_paragraph("\nNote: * p<0.1; ** p<0.05; *** p<0.01. The z (or t) statistics are reported in parentheses.")

        doc.save(filepath)
        print(f"Word document saved to: {filepath}")