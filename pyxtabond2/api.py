import pandas as pd
import numpy as np
import scipy.stats as stats
from scipy.linalg import svd
from datetime import datetime
from statsmodels.iolib.summary import Summary
from statsmodels.iolib.table import SimpleTable

from .data_utils import PanelData
from .gmm_builder import SystemGMMBuilder
from .engine import GMMEngine
from .numfac import estimate_num_factors
from .numfac import show_factor_selection

class PyXtabond2Results:
    """
    A class to store, format, and display the results of a GMM estimation.
    
    This class handles the presentation of coefficients, standard errors, 
    and diagnostic tests (Arellano-Bond AR, Sargan/Hansen, Difference-in-Sargan).
    It also provides export functionalities to LaTeX and Word.
    """
    def __init__(self, beta, se, engine, diag, x_names, dep_var, model_name, step_name, 
                 gmm_vars=None, iv_vars=None, lag_limits=None, collapse=False, orthogonal=False, id_col=str, time_col=str):
        self.beta = beta
        self.se = se
        self.engine = engine
        self.diag = diag
        self.x_names = x_names
        self.dep_var = dep_var
        self.small = engine.small
        self.model_name = model_name
        self.step_name = step_name
        
        # --- ATTRIBUTES FOR DISPLAY AND EXPORT ---
        self.gmm_vars = gmm_vars or []
        self.iv_vars = iv_vars or []
        self.lag_limits = lag_limits or (1, None)
        self.collapse = collapse
        self.orthogonal = orthogonal
        
        self.id_col = id_col
        self.time_col = time_col

        self.t_stats = self.beta / self.se
        
        if self.small:
            is_robust_or_2step = (step_name == "two-step") or hasattr(engine, 'V2_robust')
            self.df_resid = (engine.N_groups - 1) if is_robust_or_2step else (engine.n_obs - engine.k_vars)
            
            self.p_values = 2.0 * (1.0 - stats.t.cdf(np.abs(self.t_stats), self.df_resid))
            self.ci_crit = stats.t.ppf(0.975, self.df_resid)
        else:
            self.p_values = 2.0 * (1.0 - stats.norm.cdf(np.abs(self.t_stats)))
            self.ci_crit = stats.norm.ppf(0.975)

        self.ci_lower = self.beta - self.ci_crit * self.se
        self.ci_upper = self.beta + self.ci_crit * self.se

    def summary(self):
        """
        Displays the estimation results with clear, aligned typography.
        Replicates the visual structure of Stata's output.
        """
        from statsmodels.iolib.summary import Summary
        from datetime import datetime
        smry = Summary()

        # --- PREPARE GLOBAL STATISTICS ---
        if self.model_name == "system GMM":
            n_obs = int(np.sum(self.engine.is_level))
        else:
            n_obs = self.engine.n_obs
            
        n_groups = self.engine.N_groups
        n_instr = self.engine.n_instruments
        wald_stat, wald_df, wald_p, wald_type = self.diag['wald']
        is_robust_or_2step = (self.step_name == "two-step") or hasattr(self.engine, 'V2_robust')

        class MockModel:
            endog_names = self.dep_var
            exog_names = self.x_names

        class MockResults:
            def __init__(self, parent):
                self.model = MockModel() 
                self.params = parent.beta.flatten()
                self.bse = parent.se.flatten()
                self.tvalues = parent.t_stats.flatten()
                self.pvalues = parent.p_values.flatten()
                self.ci_l = parent.ci_lower.flatten()
                self.ci_u = parent.ci_upper.flatten()
            
            def conf_int(self, alpha=0.05):
                return np.column_stack((self.ci_l, self.ci_u))

        mock_res = MockResults(self)

        cov_type = "Standard"
        if hasattr(self.engine, 'V2_robust'):
            cov_type = "Windmeijer" if self.step_name == "two-step" else "Robust"

        gmm_str = ", ".join(self.gmm_vars) if self.gmm_vars else "None"
        iv_str = ", ".join(self.iv_vars) if self.iv_vars else "None"
        lag_max = self.lag_limits[1] if self.lag_limits[1] is not None else 'Max'
        lags_str = f"({self.lag_limits[0]}, {lag_max})"

        now = datetime.now()
        
        if wald_type == 'F':
            is_const = '_cons' in self.x_names
            if is_robust_or_2step:
                df_resid = self.engine.N_groups - (1 if is_const else 0)
            else:
                df_resid = self.engine.n_obs - self.engine.k_vars
                
            stat_name = f'F-statistic ({int(wald_df)}, {int(df_resid)})'
        else:
            stat_name = f'Wald chi2({int(wald_df)})'

        prob_format = f"{wald_p:.4e}" if wald_p < 0.001 else f"{wald_p:.4f}"

        top_left = [
            ('Dep. Variable:', [self.dep_var]),
            ('Model:', [f"GMM ({self.model_name.capitalize()})"]),
            ('Method:', [self.step_name.title()]),
            ('Date:', [now.strftime('%a, %d %b %Y')]),
            ('Time:', [now.strftime('%H:%M:%S')]),
            ('No. Observations:', [str(n_obs)]),
            ('No. Groups:', [str(n_groups)]),
            ('Obs per group:', [f"{n_obs/n_groups:.2f}"]),
            ('No. Instruments:', [str(n_instr)]),
            ('Covariance Type:', [cov_type])
        ]
        
        top_right = [
            (stat_name + ':', [f"{wald_stat:.2f}"]),
            ('Prob (Stat):', [prob_format]),
            ('Group Variable:', [str(self.id_col)]),
            ('Time variable:', [str(self.time_col)]),
            ('GMM Variables:', [gmm_str]),
            ('IV Variables:', [iv_str]),
            ('Lag Limits:', [lags_str]),
            ('Collapse:', ["Yes" if self.collapse else "No"]),
            ('Transform (FOD):', ["Yes" if self.orthogonal else "No (FD)"]),
            ('Sample Correction:', ["Yes" if self.small else "No"])
        ]

        # 1. Generate ONLY the header first to capture its exact width
        smry.add_table_2cols(mock_res, gleft=top_left, gright=top_right, title="PyXtabond2 Estimation Results")
        header_text = smry.as_text()
        header_w = max(len(line.rstrip()) for line in header_text.split('\n'))
        
        # 2. Add the coefficients (free to stretch)
        smry.add_table_params(mock_res, xname=self.x_names, yname=self.dep_var, use_t=self.small)
        final_lines = [line.rstrip() for line in smry.as_text().split('\n')]

        # --- PREPARE RAW DIAGNOSTIC DATA ---
        diag_data = [
            ["Diagnostic Test", "Statistic", "P-value"],
            ["Arellano-Bond AR(1)", f"z = {self.diag['ar1'][0]:.2f}", f"Pr > z = {self.diag['ar1'][1]:.3f}"],
            ["Arellano-Bond AR(2)", f"z = {self.diag['ar2'][0]:.2f}", f"Pr > z = {self.diag['ar2'][1]:.3f}"],
            ["Sargan test of overid.", f"chi2({int(self.diag['df'])}) = {self.diag['sargan'][0]:.2f}", f"Pr > chi2 = {self.diag['sargan'][1]:.3f}"]
        ]
        
        if not np.isnan(self.diag['hansen'][0]):
            diag_data.append(["Hansen test of overid.", f"chi2({int(self.diag['df'])}) = {self.diag['hansen'][0]:.2f}", f"Pr > chi2 = {self.diag['hansen'][1]:.3f}"])
        elif hasattr(self.engine, 'hansen_1step_robust'):
            h_stat, h_p = self.engine.hansen_1step_robust
            diag_data.append(["Hansen test of overid.", f"chi2({int(self.diag['df'])}) = {h_stat:.2f}", f"Pr > chi2 = {h_p:.3f}"])

        diff_tests = self.diag.get('diff_sargan', [])
        diff_data = [["Test", "Statistic", "P-value"]]
        if diff_tests:
            for res in diff_tests:
                diff_data.append([f"{res['name']} (Excl. group)", f"chi2({int(res['df_rest'])}) = {res['stat_rest']:.2f}", f"{res['p_rest']:.3f}"])
                diff_data.append(["  Difference", f"chi2({int(res['diff_df'])}) = {res['diff_stat']:.2f}", f"{res['diff_p']:.3f}"])

        # --- ALIGNMENT LOGIC WITH GENEROUS SPACING ---
        def build_matched_table(data_rows, title, target_w):
            c3_max = max(len(str(r[2])) for r in data_rows)
            c2_max = max(len(str(r[1])) for r in data_rows)
            
            col3_w = max(20, c3_max + 1)
            col2_w = max(20, c2_max + 2)
            
            col1_w = target_w - col2_w - col3_w
            
            c1_max = max(len(str(r[0])) for r in data_rows)
            if col1_w < c1_max:
                col1_w = c1_max
                rem = target_w - col1_w
                col3_w = max(10, rem // 2)
                col2_w = rem - col3_w
                
            lines = []
            lines.append("")
            lines.append(title.center(target_w))
            lines.append('=' * target_w)
            
            lines.append(f"{str(data_rows[0][0]):<{col1_w}}{str(data_rows[0][1]):>{col2_w}}{str(data_rows[0][2]):>{col3_w}}")
            lines.append('-' * target_w)
            
            for row in data_rows[1:]:
                lines.append(f"{str(row[0]):<{col1_w}}{str(row[1]):>{col2_w}}{str(row[2]):>{col3_w}}")
                
            return lines

        # 3. Inject tables 
        final_lines.extend(build_matched_table(diag_data, "Diagnostic Tests", header_w))
        final_lines.append('=' * header_w)
        
        if diff_tests:
            test_type = "Hansen" if hasattr(self.engine, 'V2_robust') or self.step_name == "two-step" else "Sargan"
            final_lines.extend(build_matched_table(diff_data, f"Difference-in-{test_type} Tests", header_w))
            final_lines.append('=' * header_w)

        # Terminal output
        print("\n".join(final_lines))
        
    # ==========================================
    # --- EXPORT METHODS (LATEX & WORD) ---
    # ==========================================

    def _get_results_df(self):
        """Utility method to structure coefficients in a DataFrame."""
        stat_name = 't' if self.small else 'z'
        df = pd.DataFrame({
            'Coefficient': self.beta.flatten(),
            'Std. Error': self.se.flatten(),
            stat_name: self.t_stats.flatten(),
            f'P>|{stat_name}|': self.p_values.flatten(),
            'CI Lower': self.ci_lower.flatten(),
            'CI Upper': self.ci_upper.flatten()
        }, index=self.x_names)
        return df

    def _get_diagnostics_df(self):
        """Utility method to structure diagnostic tests in a DataFrame."""
        diag_data = [
            ["Arellano-Bond AR(1)", f"{self.diag['ar1'][0]:.2f}", f"{self.diag['ar1'][1]:.3f}"],
            ["Arellano-Bond AR(2)", f"{self.diag['ar2'][0]:.2f}", f"{self.diag['ar2'][1]:.3f}"],
            ["Sargan test", f"{self.diag['sargan'][0]:.2f}", f"{self.diag['sargan'][1]:.3f}"]
        ]
        
        if not np.isnan(self.diag['hansen'][0]):
            diag_data.append(["Hansen test", f"{self.diag['hansen'][0]:.2f}", f"{self.diag['hansen'][1]:.3f}"])
        elif hasattr(self.engine, 'hansen_1step_robust'):
            h_stat, h_p = self.engine.hansen_1step_robust
            diag_data.append(["Hansen test", f"{h_stat:.2f}", f"{h_p:.3f}"])

        diff_tests = self.diag.get('diff_sargan', [])
        if diff_tests:
            for res in diff_tests:
                diag_data.append([f"Diff-in-Sargan: {res['name']}", f"{res['diff_stat']:.2f}", f"{res['diff_p']:.3f}"])

        return pd.DataFrame(diag_data, columns=["Test", "Statistic", "P-value"])

    def to_latex(self, filepath=None, full_output=False):
        """
        Generates a LaTeX export of the estimation results.
        
        Parameters
        ----------
        filepath : str, optional
            If provided, writes the LaTeX output to this file path.
        full_output : bool, optional
            If True, exports the raw console output in a verbatim environment 
            (ideal for appendices). If False, generates a clean academic table 
            using booktabs style. Default is False.

        Returns
        -------
        str
            The generated LaTeX string.
        """
        import io
        import sys

        if full_output:
            # --- CONSOLE CAPTURE ---
            capture = io.StringIO()
            sys.stdout = capture  # Redirect print statements to our variable
            self.summary()        # Execute summary() silently
            sys.stdout = sys.__stdout__  # Restore normal output
            full_text = capture.getvalue() # Retrieve intercepted text
            
            latex_str = "% Exact raw output from PyXtabond2 estimation\n"
            latex_str += "\\begin{verbatim}\n"
            latex_str += full_text
            latex_str += "\\end{verbatim}\n"
            
        else:
            # --- OPTION 2: ACADEMIC TABLE ---
            res_df = self._get_results_df()
            diag_df = self._get_diagnostics_df()

            # Table header
            latex_str = "\\begin{table}[htbp]\n\\centering\n\\caption{GMM Estimation Results (" + self.model_name + ")}\n"
            
            # Formatting with booktabs (requires \usepackage{booktabs} in main.tex)
            try:
                # For Pandas >= 1.3.0
                latex_str += res_df.style.format(precision=4).to_latex(hrules=True)
            except AttributeError:
                # Fallback for older Pandas versions
                latex_str += res_df.to_latex(float_format="%.4f", escape=False)
                
            latex_str += "\n\\vspace{0.5cm}\n"
            
            # Adding diagnostic tests
            try:
                 latex_str += diag_df.style.hide(axis="index").to_latex(hrules=True)
            except AttributeError:
                 latex_str += diag_df.to_latex(index=False, escape=False)

            latex_str += "\\end{table}"

        # --- COMMON SAVE LOGIC ---
        if filepath:
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(latex_str)
            mode_str = "raw (verbatim)" if full_output else "academic (booktabs)"
            print(f"LaTeX output {mode_str} successfully saved to: {filepath}")

        return latex_str

    def to_word(self, filepath="xtabond2_results.docx", full_output=False):
        """
        Exports the estimation results to a Word document (.docx).
        
        Requires the 'python-docx' library.

        Parameters
        ----------
        filepath : str, optional
            The file path where the document will be saved. Default is "xtabond2_results.docx".
        full_output : bool, optional
            If True, exports the raw console output in a monospace font. 
            If False, generates minimalist academic tables. Default is False.
        """
        import io
        import sys
        
        try:
            from docx import Document
            from docx.shared import Pt, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            print("Error: The 'python-docx' module is required.")
            return

        doc = Document()

        if full_output:
            doc.add_heading(f"Estimation Results (Raw Output): {self.model_name.capitalize()} GMM", level=0)
            
            for section in doc.sections:
                section.left_margin = Cm(1.5)
                section.right_margin = Cm(1.5)

            # --- CONSOLE CAPTURE ---
            capture = io.StringIO()
            sys.stdout = capture
            self.summary()
            sys.stdout = sys.__stdout__
            full_text = capture.getvalue()

            p = doc.add_paragraph()
            run = p.add_run(full_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(8)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        else:
            # --- OPTION 2: ACADEMIC TABLES ---
            doc.add_heading(f"Estimation Results: {self.model_name.capitalize()} GMM", level=0)

            # 1. Global metadata
            if self.model_name == "system GMM":
                import numpy as np
                n_obs = int(np.sum(self.engine.is_level))
            else:
                n_obs = self.engine.n_obs

            doc.add_paragraph(f"Dependent Variable: {self.dep_var}")
            doc.add_paragraph(f"Method: {self.step_name.title()}")
            doc.add_paragraph(f"Observations: {n_obs}")
            doc.add_paragraph(f"Number of groups: {self.engine.N_groups}")
            doc.add_paragraph(f"Number of instruments: {self.engine.n_instruments}")

            # 2. Coefficients table
            doc.add_heading("Coefficients", level=1)
            res_df = self._get_results_df()
            
            table = doc.add_table(rows=1, cols=len(res_df.columns) + 1)
            table.style = 'Table Grid'
            
            # Headers
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Variable"
            for i, col_name in enumerate(res_df.columns):
                hdr_cells[i+1].text = col_name

            # Filling variables
            for index, row in res_df.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(index)
                for i, val in enumerate(row):
                    row_cells[i+1].text = f"{val:.4f}" if isinstance(val, (int, float)) else str(val)

            # 3. Diagnostic table
            doc.add_heading("Diagnostic Tests", level=1)
            diag_df = self._get_diagnostics_df()
            
            table_diag = doc.add_table(rows=1, cols=len(diag_df.columns))
            table_diag.style = 'Table Grid'

            hdr_cells_diag = table_diag.rows[0].cells
            for i, col_name in enumerate(diag_df.columns):
                hdr_cells_diag[i].text = col_name

            for _, row in diag_df.iterrows():
                row_cells = table_diag.add_row().cells
                for i, val in enumerate(row):
                    row_cells[i].text = str(val)

        # --- COMMON SAVE LOGIC ---
        doc.save(filepath)
        mode_str = "raw" if full_output else "academic"
        print(f"Word output ({mode_str}) successfully saved to: {filepath}")
        
class PyXtabond2:
    """
    Main user interface for PyXtabond2, replicating the options of Stata's `xtabond2`.
    
    This class orchestrates data preparation, instrument building, and model 
    estimation for both standard Dynamic Panel GMM and PCA-GMM (Interactive Fixed Effects).

    Parameters
    ----------
    df : pd.DataFrame
        The panel dataset.
    id_col : str
        The name of the column identifying the panel groups.
    time_col : str
        The name of the column identifying time periods.
    dep_var : str
        The dependent variable.
    x_vars : list of str
        The list of strictly exogenous variables to be included as standard regressors.
    gmm_vars : list of str
        The list of endogenous/predetermined variables to be instrumented GMM-style.
    iv_vars : list of str
        The list of exogenous variables to be used as standard IV instruments.
    model_type : str, optional
        'difference' for Arellano-Bond, or 'system' for Arellano-Bover/Blundell-Bond. Default is 'difference'.
    twostep : bool, optional
        If True, estimates the two-step GMM. Default is False (one-step).
    robust : bool, optional
        If True, computes robust standard errors (Windmeijer corrected for two-step). Default is False.
    lag_limits_diff : tuple, optional
        The lag limits for GMM instruments (e.g., `(2, None)` means lags 2 to maximum).
    collapse : bool, optional
        If True, limits instrument proliferation by collapsing the instrument matrix.
    orthogonal : bool, optional
        If True, uses Forward Orthogonal Deviations (FOD) instead of first differences.
    small : bool, optional
        If True, applies small-sample degree-of-freedom corrections to covariance matrices and test statistics.
    r : int or str, optional
        The number of interactive unobserved factors for PCA-GMM. If 'auto', estimates the optimal number of factors. Default is 0.
    r_max : int, optional
        Maximum number of factors to test if `r='auto'`. Default is 5.
    ife_max_iter : int, optional
        Maximum number of iterations for the PCA-GMM convergence loop. Default is 30.
    ife_tol : float, optional
        Convergence tolerance limit for PCA-GMM. Default is 1e-4.

    Raises
    ------
    ValueError
        If `model_type` is not 'difference' or 'system'.
    """
    def __init__(self, df: pd.DataFrame, id_col: str, time_col: str, 
                 dep_var: str, x_vars: list, gmm_vars: list, iv_vars: list,
                 model_type: str = 'difference',  # 'difference' or 'system'
                 twostep: bool = False,        # False = One-step, True = Two-step
                 robust: bool = False,         # Robust variance matrix
                 lag_limits_diff: tuple = (2, None), 
                 collapse: bool = False, 
                 orthogonal: bool = False, 
                 small: bool = False,
                 r: int = 0,
                 r_max: int = 5,
                 ife_max_iter: int = 30, 
                 ife_tol: float = 1e-4):
        
        if model_type not in ['difference', 'system']:
            raise ValueError("model_type must be 'difference' or 'system'")
            
        self.df = df.copy()
        self.id_col = id_col
        self.time_col = time_col
        self.dep_var = dep_var
        self.x_vars = x_vars.copy()
        self.gmm_vars = gmm_vars
        self.iv_vars = iv_vars
        self.model_type = model_type
        self.twostep = twostep
        self.robust = robust
        self.lag_limits_diff = lag_limits_diff
        self.collapse = collapse
        self.orthogonal = orthogonal
        self.small = small
        self.r = r
        self.r_max = r_max
        self.ife_max_iter = ife_max_iter
        self.ife_tol = ife_tol

    def fit(self):
        """
        Main method: Routes between Classic GMM and GMM with Interactive Fixed Effects (IFE).
        
        Returns
        -------
        PyXtabond2Results
            The structured estimation results.
        """
        if self.r == 0:
            # Classic GMM (Direct call to the base engine)
            return self._fit_base(self.df, self.dep_var)
        else:
            # GMM with Interactive Fixed Effects (PCA-GMM)
            return self._fit_ife()
        
    def _fit_ife(self):
        """
        Under-the-hood iterative engine for PCA-GMM.
        Handles the extraction and purging of unobserved common factors.
        """
        print(f"\n--- STARTING PCA-GMM ESTIMATION ---")
        df_work = self.df.copy()
        y_tilde_col = f"{self.dep_var}_tilde"

        # ========================================================
        # 1. SAFEGUARD: STEP 0 IN PURE GMM
        # ========================================================
        choix_r_utilisateur = self.r  # Save the user argument ("auto" or integer)
        self.r = 0                    # Force r=0 so the engine does not penalize degrees of freedom yet
        
        # Naive estimation (Biased)
        res_iter = self._fit_base(df_work, self.dep_var)
        beta_old = res_iter.beta.flatten()
        
        # ========================================================
        # 2. AUTO-DETECTION (Strict Sandbox)
        # ========================================================
        if str(choix_r_utilisateur).lower() == 'auto':
            # Recreate the initial residuals purely in memory
            beta_dict = dict(zip(res_iter.x_names, res_iter.beta.flatten()))
            resid_temp = df_work[self.dep_var].copy()
            
            for x_name in self.x_vars:
                if x_name in beta_dict:
                    resid_temp -= beta_dict[x_name] * df_work[x_name]
            if '_cons' in beta_dict:
                resid_temp -= beta_dict['_cons']
                
            # Pivot to (N x T) matrix
            resid_matrix = pd.DataFrame({'id': df_work[self.id_col], 'time': df_work[self.time_col], 'e': resid_temp})
            E_mat = resid_matrix.pivot(index='id', columns='time', values='e').values
            E_mat = np.nan_to_num(E_mat, nan=np.nanmean(E_mat))
            
            self.fac_results = estimate_num_factors(E_mat, kmax=getattr(self, 'r_max', 5))
            self.r = self.fac_results['best_er'] # Continue using ER metric for the rest of the algorithm
        else:
            self.r = int(choix_r_utilisateur)

        print(f">> PCA-GMM algorithm configured with r = {self.r}")

        for iteration in range(1, self.ife_max_iter + 1):
            # 1. Compute generic level residuals (y_it - X_it * beta)
            beta_dict = dict(zip(res_iter.x_names, res_iter.beta.flatten()))
            df_work['resid_ife'] = df_work[self.dep_var]
            
            for x_name in self.x_vars:
                if x_name in beta_dict:
                    df_work['resid_ife'] -= beta_dict[x_name] * df_work[x_name]
            if '_cons' in beta_dict:
                df_work['resid_ife'] -= beta_dict['_cons']

            # 2. Factor Extraction (SVD)
            resid_matrix = df_work.pivot(index=self.id_col, columns=self.time_col, values='resid_ife')
            ids = resid_matrix.index
            times = resid_matrix.columns
            E_mat = resid_matrix.values
            E_mat = np.nan_to_num(E_mat, nan=np.nanmean(E_mat)) # Safety against NaNs

            U, S, Vt = svd(E_mat, full_matrices=False)
            N_indiv = len(ids)
            Lambda = U[:, :self.r] * np.sqrt(N_indiv)
            F = (Vt[:self.r, :].T * S[:self.r]) / np.sqrt(N_indiv)
            
            # Reconstruct the factorial structure
            IFE_mat = Lambda @ F.T

            # 3. Purge the dependent variable
            df_ife = pd.DataFrame(IFE_mat, index=ids, columns=times)
            df_ife_long = df_ife.reset_index().melt(id_vars=self.id_col, var_name=self.time_col, value_name='ife_component')

            if 'ife_component' in df_work.columns:
                df_work = df_work.drop(columns=['ife_component'])
            df_work = df_work.merge(df_ife_long, on=[self.id_col, self.time_col], how='left')

            # New target variable for the next iteration
            df_work[y_tilde_col] = df_work[self.dep_var] - df_work['ife_component']

            # 4. GMM step on purged data
            # Note: We pass y_tilde as the target, 
            # but instruments will be built on raw y (since gmm_vars=['y'])
            res_iter = self._fit_base(df_work, y_tilde_col)
            beta_new = res_iter.beta.flatten()

            # 5. Evaluate convergence
            diff = np.sum(np.abs(beta_new - beta_old))
            print(f"Iteration {iteration:02d} | Convergence delta: {diff:.6f}")

            if diff < self.ife_tol:
                print(f"✅ PCA-GMM convergence achieved at iteration {iteration}!")
                break
                
            beta_old = beta_new

        if diff >= self.ife_tol:
            print(f"⚠️ Convergence not achieved after {self.ife_max_iter} iterations.")

        # Cosmetic trick: Restore the true variable name for the final table
        res_iter.dep_var = self.dep_var 
        return res_iter
    
    def plot_factor_selection(self, mode='graph'):
        """
        Displays the table or plot for factor selection criteria.
        Only functions if the model was estimated with r='auto'.
        """
        if not hasattr(self, 'fac_results'):
            print("⚠️ Cannot display criteria: PCA-GMM estimation with r='auto' has not been executed yet.")
            return
            
        return show_factor_selection(self.fac_results, mode=mode)
    
    def _fit_base(self, df_custom, dep_var_custom):
        """Prepares data, builds instruments according to model type, and estimates."""
        panel = PanelData(df_custom, self.id_col, self.time_col)
        df_clean = panel.data
        t_min = df_clean.index.get_level_values(self.time_col).min()
        t_max = df_clean.index.get_level_values(self.time_col).max()
        T_span = t_max - t_min + 1
        
        df_clean['_cons'] = 1.0
        df_clean['D__cons'] = 0.0
        if self.orthogonal: df_clean['FOD__cons'] = 0.0
        
        trans_prefix = 'FOD_' if self.orthogonal else 'D_'
        all_vars = set([dep_var_custom] + self.x_vars + self.gmm_vars + self.iv_vars)
        if '_cons' in all_vars: all_vars.remove('_cons')
        
        for v in all_vars:
            if trans_prefix + v not in df_clean.columns:
                df_clean[trans_prefix + v] = panel.get_fod(v) if self.orthogonal else panel.get_first_difference(v)
            if 'D_' + v not in df_clean.columns: 
                df_clean['D_' + v] = panel.get_first_difference(v)

        sys_builder = SystemGMMBuilder(panel)
        Z_list = []
        
        # 1. Building Instrument Matrices
        for gvar in self.gmm_vars:
            Z_full = sys_builder.build_system_instruments(gvar, self.lag_limits_diff, self.collapse)
            if self.model_type == 'difference':
                # The number of "level" columns to truncate depends on the collapse option!
                cols_to_drop = 1 if self.collapse else (T_span - 1)
                
                # Using Python [:-n] syntax to cleanly remove the last 'n' columns
                Z_list.append(Z_full[:, :-cols_to_drop]) 
            else:
                Z_list.append(Z_full)
                
        for ivar in self.iv_vars:
            include_lvl = (self.model_type == 'system')
            Z_iv = sys_builder.build_iv_instruments(ivar, include_in_diff=True, include_in_level=include_lvl, orthogonal=self.orthogonal)
            
            # --- XTABOND2 SECRET FOR STANDARD IVs ---
            # In System GMM, iv() combines the difference and the level into A SINGLE column!
            if include_lvl and Z_iv.shape[1] > 1:
                Z_iv = np.sum(Z_iv, axis=1, keepdims=True)
                
            Z_list.append(Z_iv)
                
        if '_cons' not in self.iv_vars and self.model_type == 'system':
            Z_list.append(sys_builder.build_iv_instruments('_cons', include_in_diff=False, include_in_level=True, orthogonal=self.orthogonal))
            
        Z_sys_full = np.hstack(Z_list)

        # 2. Stacking Y and X according to chosen model
        Y_stacked, X_stacked, Z_stacked = [], [], []
        Y_lvl_list, X_lvl_list = [], [] # <-- NEW
        group_ids_list, is_level_list = [], []
        t_idx_list = [] # NEW: To track exact time
        
        # In difference GMM, there is no constant.
        x_names_final = self.x_vars + ['_cons'] if self.model_type == 'system' else self.x_vars
        unique_groups = df_clean.index.get_level_values(0).unique()
        z_idx = 0

        for g in unique_groups:
            mask_g = (df_clean.index.get_level_values(0) == g)
            df_g = df_clean[mask_g]
            
            mask_diff = df_g.index.get_level_values(self.time_col) >= (t_min + 1)
            y_diff = df_g.loc[mask_diff, trans_prefix + dep_var_custom].values.reshape(-1, 1)
            
            # NEW: Extract relative time indices
            t_diff = df_g.index.get_level_values(self.time_col)[mask_diff].values - t_min
            
            if self.model_type == 'difference':
                # --- NEW: Extracting exact Levels for the AR Test ---
                y_lvl_g = df_g.loc[mask_diff, dep_var_custom].values.reshape(-1, 1)
                X_lvl_cols = [df_g.loc[mask_diff, x].values.reshape(-1, 1) for x in self.x_vars]
                X_lvl_g = np.hstack(X_lvl_cols) if X_lvl_cols else np.zeros((len(y_lvl_g), 0))
                Y_lvl_list.append(y_lvl_g)
                X_lvl_list.append(X_lvl_g)
                
            else:
                y_lvl_g = df_g[dep_var_custom].values.reshape(-1, 1)
                X_lvl_cols = [df_g[x].values.reshape(-1, 1) for x in self.x_vars]
                X_lvl_g = np.hstack(X_lvl_cols) if X_lvl_cols else np.zeros((len(y_lvl_g), 0))
                Y_lvl_list.append(y_lvl_g)
                X_lvl_list.append(X_lvl_g)
            
            if self.model_type == 'system':
                X_diff_cols = [df_g.loc[mask_diff, trans_prefix + x].values.reshape(-1, 1) for x in self.x_vars]
                X_diff = np.hstack(X_diff_cols + [np.zeros((len(y_diff), 1))])
                
                mask_level = df_g.index.get_level_values(self.time_col) >= (t_min + 1)
                y_level = df_g.loc[mask_level, dep_var_custom].values.reshape(-1, 1)
                X_lvl_cols = [df_g.loc[mask_level, x].values.reshape(-1, 1) for x in self.x_vars]
                X_level = np.hstack(X_lvl_cols + [np.ones((len(y_level), 1))])
                
                # NEW: Time index for levels (offset by T_span)
                t_level = df_g.index.get_level_values(self.time_col)[mask_level].values - t_min
                
                Z_g_diff = Z_sys_full[z_idx : z_idx + T_span][mask_diff]
                Z_g_level = Z_sys_full[z_idx + T_span : z_idx + 2*T_span][mask_level]
                z_idx += 2 * T_span
                
                Y_stacked.append(np.vstack([y_diff, y_level]))
                X_stacked.append(np.vstack([X_diff, X_level]))
                Z_stacked.append(np.vstack([Z_g_diff, Z_g_level]))
                
                n_d, n_l = len(y_diff), len(y_level)
                group_ids_list.extend([g] * (n_d + n_l))
                is_level_list.extend([False] * n_d + [True] * n_l)
                
                # NEW: Append to global list
                t_idx_list.extend(t_diff)
                t_idx_list.extend(t_level + T_span)
            else:
                # Pure Difference Model
                X_diff_cols = [df_g.loc[mask_diff, trans_prefix + x].values.reshape(-1, 1) for x in self.x_vars]
                X_diff = np.hstack(X_diff_cols)
                
                Z_g_diff = Z_sys_full[z_idx : z_idx + T_span][mask_diff]
                z_idx += 2 * T_span
                
                Y_stacked.append(y_diff)
                X_stacked.append(X_diff)
                Z_stacked.append(Z_g_diff)
                
                n_d = len(y_diff)
                group_ids_list.extend([g] * n_d)
                is_level_list.extend([False] * n_d)
                t_idx_list.extend(t_diff)

        Y_sys = np.vstack(Y_stacked)
        X_sys = np.vstack(X_stacked)
        Z_sys = np.vstack(Z_stacked)
        
        group_ids = np.array(group_ids_list)
        is_level = np.array(is_level_list)
        t_index = np.array(t_idx_list) # NEW
        
        # 3. Launching the Engine with correct options
        engine = GMMEngine(
            Y_sys, X_sys, Z_sys, 
            group_ids=group_ids, is_level=is_level,
            small=self.small, orthogonal=self.orthogonal,
            t_index=t_index, T_span=T_span,
            y_lvl=Y_lvl_list, X_lvl=X_lvl_list,
            r=self.r
        )
        
        if self.twostep:
            if self.robust:
                beta, se, _ = engine.estimate_two_step_robust()
            else:
                # Non-robust Two-step (Rarely used, but Stata allows it)
                engine.estimate_two_step_robust() # Called to populate beta2
                beta = engine.beta2
                se = np.sqrt(np.diag(engine.V2)).reshape(-1, 1)
                delattr(engine, 'V2_robust')
        else:
            # One-step
            beta = engine.estimate_one_step()
            if self.robust:
                # Calculate One-Step Robust Variance
                S1_robust = np.zeros((engine.n_instruments, engine.n_instruments))
                for g in np.unique(engine.group_ids):
                    mask = (engine.group_ids == g)
                    Z_g = engine.Z[mask]
                    e1_g = engine.e1[mask]
                    S1_robust += Z_g.T @ (e1_g @ e1_g.T) @ Z_g
                
                XZ = engine.X.T @ engine.Z
                meat = XZ @ engine.W1 @ S1_robust @ engine.W1 @ XZ.T
                V1_rob = engine.V1 @ meat @ engine.V1
                
                if self.small:
                    N_obs = np.sum(engine.is_level) if self.model_type == 'system' else len(engine.y)
                    qc = (N_obs / (N_obs - engine.k_vars)) * (engine.N_groups / (engine.N_groups - 1))
                    V1_rob *= qc
                se = np.sqrt(np.diag(V1_rob)).reshape(-1, 1)
                engine.V2_robust = V1_rob # Trick for Wald test to use it

                # --- ROODMAN'S TRICK (xtabond2.ado) ---
                # Silent execution of Step 2 just to compute the Hansen test
                W2_hansen = np.linalg.pinv(S1_robust)
                Zy = engine.Z.T @ engine.y
                XZ_W2_ZX = XZ @ W2_hansen @ XZ.T
                V2_hansen = np.linalg.pinv((XZ_W2_ZX + XZ_W2_ZX.T) / 2.0)
                beta2_hansen = V2_hansen @ XZ @ W2_hansen @ Zy
                e2_hansen = engine.y - engine.X @ beta2_hansen
                
                hansen_stat = (e2_hansen.T @ engine.Z @ W2_hansen @ engine.Z.T @ e2_hansen)[0, 0]
                df_tests = engine.n_instruments - engine.k_vars
                hansen_p = 1.0 - stats.chi2.cdf(hansen_stat, df_tests) if df_tests > 0 else np.nan
                engine.hansen_1step_robust = (hansen_stat, hansen_p)
            else:
                # Non-Robust One-Step (Homoscedastic)
                divisor = 1.0 if self.orthogonal else 2.0
                
                # --- STATA CORRECTION: sig2_v based exclusively on differences ---
                if self.model_type == 'system':
                    # We sum ONLY the difference residuals!
                    e1_diff = engine.e1[~engine.is_level]
                    sum_sq = np.sum(e1_diff**2)
                    # Denominator is the total number of original level observations
                    N_obs = np.sum(~engine.is_level)
                else:
                    sum_sq = np.sum(engine.e1**2)
                    N_obs = engine.n_obs
                    
                # --- MODIFIED: PCA penalty on the denominator ---
                df_pca_penalty = self.r * (engine.N_groups + T_span - self.r) if self.r > 0 else 0
                denom = (N_obs - engine.k_vars - df_pca_penalty) if self.small else N_obs
                sig2_v = sum_sq / denom / divisor
                
                # Crucial save of this exact scale for AR() tests
                engine.sig2_v_1step = sig2_v 
                
                engine.V1 = engine.V1 * sig2_v
                se = np.sqrt(np.diag(engine.V1)).reshape(-1, 1)

        # =================================================================
        # --- DIFFERENCE-IN-SARGAN / HANSEN TESTS (MODULAR) ---
        # =================================================================
        diff_sargan_results = []
        
        # Utility function to execute the test on any sub-group
        def compute_c_stat(exclude_mask, test_name):
            if not np.any(exclude_mask): return
            
            Z_rest = engine.Z[:, ~exclude_mask]
            
            # --- THE FIX IS HERE ---
            # 1. Check basic dimensionality before even instantiating
            if Z_rest.shape[1] < engine.k_vars:
                return
                
            try:
                engine_rest = GMMEngine(
                    engine.y, engine.X, Z_rest, 
                    group_ids=engine.group_ids, is_level=engine.is_level,
                    small=self.small, orthogonal=self.orthogonal,
                    t_index=getattr(engine, 't_index', None), T_span=getattr(engine, 'T_span', None),
                    y_lvl=getattr(engine, 'y_lvl', None), X_lvl=getattr(engine, 'X_lvl', None)
                )
            except ValueError as e:
                # 2. If the engine detects insufficient matrix rank, cancel test silently
                if "Modèle sous-identifié" in str(e) or "under-identified" in str(e).lower():
                    return
                raise e
            
            df_rest = engine_rest.n_instruments - engine_rest.k_vars
            if df_rest < 0: return # Impossible to test if under-identified
            
            if self.twostep or self.robust:
                S_rest = np.zeros((engine_rest.n_instruments, engine_rest.n_instruments))
                for g in np.unique(engine.group_ids):
                    mask_g = (engine.group_ids == g)
                    e1_full_g = engine.e1[mask_g]
                    Z_rest_g = engine_rest.Z[mask_g]
                    S_rest += Z_rest_g.T @ (e1_full_g @ e1_full_g.T) @ Z_rest_g
                
                # Protecting the pseudo-inverse matrix of the sub-sample
                try:
                    W2_rest = np.linalg.pinv(S_rest)
                except np.linalg.LinAlgError:
                    return
                    
                Zy_rest = engine_rest.Z.T @ engine_rest.y
                XZ_rest = engine_rest.X.T @ engine_rest.Z
                XZ_W2_ZX_rest = XZ_rest @ W2_rest @ XZ_rest.T
                V2_rest = np.linalg.pinv((XZ_W2_ZX_rest + XZ_W2_ZX_rest.T) / 2.0)
                beta2_rest = V2_rest @ XZ_rest @ W2_rest @ Zy_rest
                e2_rest = engine_rest.y - engine_rest.X @ beta2_rest
                stat_rest = (e2_rest.T @ engine_rest.Z @ W2_rest @ engine_rest.Z.T @ e2_rest)[0, 0]
            else:
                # --- C-STATISTIC : SARGAN (Homoscedastic 1-Step) ---
                engine_rest.estimate_one_step()
                
                # THE ULTIMATE SOLUTION: Extract the exact variance directly from the full model!
                sig2_full = getattr(engine, 'sig2_v_1step', 1.0)
                
                stat_rest = (engine_rest.e1.T @ engine_rest.Z @ engine_rest.W1 @ engine_rest.Z.T @ engine_rest.e1)[0, 0] / sig2_full
                
            diag_main = engine.get_diagnostics()
            if self.twostep: stat_full = diag_main['hansen'][0]
            elif self.robust: stat_full = engine.hansen_1step_robust[0]
            else: stat_full = diag_main['sargan'][0]
                
            df_full = diag_main['df']
            diff_stat = stat_full - stat_rest
            diff_df = df_full - df_rest
            diff_p = 1.0 - stats.chi2.cdf(diff_stat, diff_df) if diff_df > 0 else np.nan
            
            diff_sargan_results.append({
                'name': test_name,
                'stat_rest': stat_rest, 'df_rest': df_rest,
                'p_rest': 1.0 - stats.chi2.cdf(stat_rest, df_rest) if df_rest > 0 else np.nan,
                'diff_stat': diff_stat, 'diff_df': diff_df, 'diff_p': diff_p
            })

        # --- Test 1 : "GMM instruments for levels" (System only) ---
        if self.model_type == 'system':
            diff_mask = ~engine.is_level
            level_mask = engine.is_level
            is_level_inst = np.all(np.abs(engine.Z[diff_mask, :]) < 1e-10, axis=0)
            is_cons_inst = is_level_inst & np.all(np.abs(engine.Z[level_mask, :] - 1.0) < 1e-10, axis=0)
            is_gmm_level_inst = is_level_inst & ~is_cons_inst
            compute_c_stat(is_gmm_level_inst, 'GMM instruments for levels')

        # --- Test 2 : "iv(x)" (Standard Instruments) ---
        if self.iv_vars:
            n_iv = len(self.iv_vars)
            iv_mask = np.zeros(engine.Z.shape[1], dtype=bool)
            
            # IV instruments are always added to the end of matrix Z.
            # In system, the constant is the very last column. In difference, they are the IVs.
            if self.model_type == 'system':
                iv_mask[-(n_iv+1):-1] = True 
            else:
                iv_mask[-n_iv:] = True
                
            iv_names = " ".join(self.iv_vars)
            compute_c_stat(iv_mask, f'iv({iv_names})')

        diag = engine.get_diagnostics()
        diag['diff_sargan'] = diff_sargan_results

        m_name = "system GMM" if self.model_type == 'system' else "difference GMM"
        s_name = "two-step" if self.twostep else "one-step"

        return PyXtabond2Results(
            beta, se, engine, diag, x_names_final, dep_var_custom, m_name, s_name,
            gmm_vars=self.gmm_vars, 
            iv_vars=self.iv_vars, 
            lag_limits=self.lag_limits_diff, 
            collapse=self.collapse, 
            orthogonal=self.orthogonal,
            id_col=self.id_col,
            time_col=self.time_col
        )